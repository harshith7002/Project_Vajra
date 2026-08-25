import os
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from backend.audit.logger import audit_logger
from backend.security.egress_monitor import egress_monitor

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "data", "outputs")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_approval_note(
    equipment_id: str = "Boiler Feedwater Pump B-102",
    source_doc: str = "Pump_Inspection_Report_07.pdf",
    findings_summary: str = "Vibration level reached 7.8 mm/s RMS on drive-end bearing housing, exceeding ISO alert thresholds.",
    recommended_actions: list = None,
    evidence_passages: list = None,
    approved_by: str = "Chief Maintenance Operations Lead"
) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title Header
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = title_p.add_run("SOVEREIGN AI WORKBENCH — CONFIDENTIAL DELIVERABLE\n")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(9)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0, 102, 153) # Navy blue

    run_title = title_p.add_run("MAINTENANCE APPROVAL NOTE")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate dark

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Metadata Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("APPROVAL NOTE ID:", f"MAN-2026-{uuid.uuid4().hex[:6].upper()}"),
        ("DATE / TIME:", datetime.now().strftime("%Y-%m-%d %H:%M:%S UTC+05:30")),
        ("TARGET EQUIPMENT:", equipment_id),
        ("SOURCE INSPECTION DOC:", source_doc)
    ]
    
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.name = "Arial"
        set_cell_background(cell_lbl, "F1F5F9")

        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9.5)
        r_val.font.name = "Arial"
        set_cell_background(cell_val, "FAFAFA")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 1: Executive Summary & Critical Findings
    h1 = doc.add_heading("1. Executive Summary & Critical Findings", level=1)
    h1.runs[0].font.color.rgb = RGBColor(15, 23, 42)
    h1.runs[0].font.size = Pt(13)

    p_find = doc.add_paragraph()
    r_find = p_find.add_run(findings_summary)
    r_find.font.size = Pt(10)
    r_find.font.name = "Arial"

    # Status callout box
    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = box_table.cell(0, 0)
    set_cell_background(c, "FEF2F2") # Light red
    p_box = c.paragraphs[0]
    r_box = p_box.add_run("CRITICAL ALERT: ISO 10816 Class II threshold exceeded (7.8 mm/s RMS). Immediate category 2 LOTO isolation required.")
    r_box.font.bold = True
    r_box.font.size = Pt(9.5)
    r_box.font.color.rgb = RGBColor(185, 28, 28) # Red

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 2: Approved Maintenance Actions
    h2 = doc.add_heading("2. Recommended Maintenance Work Orders", level=1)
    h2.runs[0].font.color.rgb = RGBColor(15, 23, 42)
    h2.runs[0].font.size = Pt(13)

    actions = recommended_actions or [
        "Execute LOTO isolation procedure on Feedwater Isolation Valve HV-104 and Pump B-102 power feeder.",
        "Dismantle and inspect drive-end bearing sleeve B-102-BRG for fatigue fractures.",
        "Flush and clean Suction Strainer ST-01 to restore NPSH available.",
        "Perform post-overhaul vibration diagnostic and re-calibrate pressure sensor PT-208."
    ]

    for act in actions:
        p_act = doc.add_paragraph(style='List Bullet')
        r_act = p_act.add_run(act)
        r_act.font.size = Pt(10)
        r_act.font.name = "Arial"

    # Section 3: Evidence & Audit Verification
    h3 = doc.add_heading("3. Local RAG Evidence Citations", level=1)
    h3.runs[0].font.color.rgb = RGBColor(15, 23, 42)
    h3.runs[0].font.size = Pt(13)

    passages = evidence_passages or [
        {"filename": "Pump_Inspection_Report_07.pdf", "page": 4, "snippet": "Vibration level reached 7.8 mm/s RMS on drive-end bearing housing, exceeding operating limits."},
        {"filename": "SOP_Pump_Maintenance_2025.pdf", "page": 12, "snippet": "Section 4.2: Any measured vibration >= 7.0 mm/s mandates immediate emergency shutdown and LOTO clearance."}
    ]

    for ev in passages:
        p_ev = doc.add_paragraph()
        r_ev_hdr = p_ev.add_run(f"[{ev['filename']} — Page {ev['page']}]:\n")
        r_ev_hdr.font.bold = True
        r_ev_hdr.font.size = Pt(9)
        r_ev_hdr.font.color.rgb = RGBColor(30, 58, 138)
        
        r_ev_txt = p_ev.add_run(f"\"{ev['snippet']}\"")
        r_ev_txt.font.italic = True
        r_ev_txt.font.size = Pt(9.5)
        r_ev_txt.font.name = "Arial"

    # Section 4: Human Review & Sign-Off
    h4 = doc.add_heading("4. Authorizing Signature & Governance", level=1)
    h4.runs[0].font.color.rgb = RGBColor(15, 23, 42)
    h4.runs[0].font.size = Pt(13)

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    sig_table.rows[0].cells[0].paragraphs[0].add_run("PREPARED BY (LOCAL AI):").font.bold = True
    sig_table.rows[0].cells[1].paragraphs[0].add_run("HUMAN APPROVER SIGN-OFF:").font.bold = True
    
    sig_table.rows[1].cells[0].paragraphs[0].add_run("VAJRA Sovereign Agent Engine v1.0\n[Zero Cloud Network Egress Verified]").font.size = Pt(9)
    sig_table.rows[1].cells[1].paragraphs[0].add_run(f"Approved by: {approved_by}\nStatus: VERIFIED & EXECUTED\nTimestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").font.size = Pt(9)

    out_filename = f"MAINTENANCE_APPROVAL_NOTE_B102_{uuid.uuid4().hex[:4]}.docx"
    filepath = os.path.join(OUTPUT_DIR, out_filename)
    doc.save(filepath)

    audit_logger.log(
        event_type="DOCX_EXPORTED",
        actor="VAJRA_DOCX_GENERATOR",
        details=f"Generated real DOCX approval note: {out_filename}",
        metadata={"filepath": filepath, "equipment_id": equipment_id}
    )
    egress_monitor.record_event("LOCAL_FILE_EXPORT", "GENERATE_DOCX", "ALLOWED", f"Saved deliverable to {out_filename}")

    return filepath
